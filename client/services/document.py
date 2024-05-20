import os
import time
from typing import List, Tuple
import io
import re
import requests  # type: ignore # noqa: F401
from dotenv import load_dotenv
from fastapi import UploadFile
from PyPDF2 import PdfReader
import docx

from .key import encrypt_document, encrypt_ndarray, ctxt_to_str
from ..models.document import Document, PyCDocumentDto
from .embedding import embed_sentence
from .session import get_user_id


async def read_file(file: UploadFile) -> Tuple[str, int]:
    contents = await file.read()
    file_size = os.fstat(file.file.fileno()).st_size
    file_extension = file.filename.split(".")[-1].lower()

    if file_extension == "pdf":
        pdf_reader = PdfReader(io.BytesIO(contents))
        str_content = ""
        for page in pdf_reader.pages:
            str_content += page.extract_text()

    elif file_extension == "docx":
        doc = docx.Document(io.BytesIO(contents))
        str_content = ""
        for para in doc.paragraphs:
            str_content += para.text

    else:
        str_content = contents.decode("utf-8")

    return str_content, file_size


def split_content(str_content: str, chunk_size: int = 300) -> List[Document]:
    docs = []
    index = 0
    start = 0
    pattern = r"[ \n\.]"  # 공백 문자, 개행 문자, 마침표를 감지하는 정규식 패턴

    while start < len(str_content):
        end = min(start + chunk_size, len(str_content))
        chunk = str_content[start:end]

        # 공백 문자, 개행 문자, 마침표를 찾아 문자열을 나눕니다.
        match = re.search(pattern, chunk[chunk_size:])
        if match:
            end = start + match.start() + chunk_size

        chunk = str_content[start:end].strip()
        doc = Document(index=index, document=chunk)
        docs.append(doc)
        index += 1
        start = end

    return docs


def embed_documents(documents: List[Document]) -> Tuple[List[Document], float]:
    embed_times = []

    for document in documents:
        doc_start_time = time.time()
        emb = embed_sentence(sentence=document.document)
        document.embedding = emb
        doc_end_time = time.time()
        doc_embed_time = doc_end_time - doc_start_time
        embed_times.append(doc_embed_time)
        print(f"embedding completed, ({document.index + 1}/{len(documents)})")

    avg_embed_time = sum(embed_times) / len(embed_times) if embed_times else 0.0

    return documents, round(avg_embed_time, 3)


def encrypt_documents(documents: List[Document]) -> Tuple[List[PyCDocumentDto], float]:
    encrypted_documents = []
    encrypt_times = []

    for document in documents:
        doc_start_time = time.time()
        enc_doc = encrypt_document(document.document)
        enc_emb = encrypt_ndarray(document.embedding)
        doc_end_time = time.time()
        doc_encrypt_time = doc_end_time - doc_start_time
        encrypt_times.append(doc_encrypt_time)

        encrypted_document = PyCDocumentDto(
            index=document.index,
            document=ctxt_to_str(enc_doc),
            embedding=ctxt_to_str(enc_emb),
        )

        encrypted_documents.append(encrypted_document)

    avg_encrypt_time = sum(encrypt_times) / len(encrypt_times) if encrypt_times else 0.0

    return encrypted_documents, round(avg_encrypt_time, 3)


def send_documents(uri: str, encrypted_documents: List[PyCDocumentDto]):
    load_dotenv()

    server_url = os.getenv("SERVER_URL") or "EMPTY"

    user_id = get_user_id()
    headers = {"origin": user_id}

    json = [doc.to_dict() for doc in encrypted_documents]
    response = requests.post(server_url + uri, headers=headers, json=json, timeout=120)

    return response
